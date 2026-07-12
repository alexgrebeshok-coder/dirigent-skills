#!/usr/bin/env python3
"""Post-process DOCX to a fixed Russian business document standard (Times New Roman, black, RU)."""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    print("Install: python3 -m pip install python-docx", file=sys.stderr)
    sys.exit(1)

# Text replacements (order matters — longer first)
TEXT_REPLACEMENTS = [
    ("Make-up right", "Право на последующую выборку"),
    ("make-up right", "право на последующую выборку"),
    ("Take-or-pay", "Условие принятия либо оплаты"),
    ("take-or-pay", "условие принятия либо оплаты"),
    ("Take or pay", "Условие принятия либо оплаты"),
    ("deliver-or-pay", "условие принятия либо оплаты"),
    ("fallback", "запасной переговорный вариант"),
    ("Fallback", "Запасной переговорный вариант"),
    ("₽/m³", "₽/м³"),
    ("₽/m3", "₽/м³"),
    ("m³", "м³"),  # careful: only if latin m — applied in context
]

HOMOGLYPH_FIXES = [
    ("m³", "м³"),
    ("/m³", "/м³"),
]


def replace_text(text: str, external: bool = True) -> str:
    for old, new in TEXT_REPLACEMENTS:
        text = text.replace(old, new)
    for old, new in HOMOGLYPH_FIXES:
        # avoid breaking URLs
        if "http" in text and old in text:
            continue
        text = text.replace(old, new)
    if external:
        text = re.sub(r"\blegalization\b", "оформление в товарный режим", text, flags=re.I)
    return text


def normalize_run(run, size_pt: int = 12) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)


def heading_size(style_name: str, flat: bool = False) -> int:
    if flat:
        return 12
    if "Heading 1" in style_name or style_name == "Title":
        return 16
    if "Heading 2" in style_name:
        return 14
    if "Heading 3" in style_name:
        return 13
    return 12


def process_paragraph(para, default_size: int = 12, flat: bool = False) -> None:
    size = heading_size(getattr(para.style, "name", "") or "", flat=flat)
    if size == 12:
        size = default_size
    if flat and para.text.strip():
        for run in para.runs:
            run.bold = True if "Heading" in (getattr(para.style, "name", "") or "") or getattr(para.style, "name", "") == "Title" else run.bold
    new_text = replace_text(para.text)
    if new_text != para.text:
        para.text = new_text
    for run in para.runs:
        normalize_run(run, size)


def process_docx(path: Path, default_size: int = 12, flat: bool = False) -> None:
    doc = Document(str(path))
    for para in doc.paragraphs:
        process_paragraph(para, default_size, flat=flat)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para, default_size, flat=flat)
    doc.save(str(path))


def scan_bad_patterns(path: Path) -> list[str]:
    import zipfile

    issues: list[str] = []
    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8", errors="replace")
    text = re.sub(r"<[^>]+>", "", xml)
    patterns = [
        (r"Table of Contents", "англ. оглавление"),
        (r"Служебная пометка", "служебная пометка"),
        (r"не для отправки", "метка «не для отправки»"),
        (r"take-or-pay|make-up|fallback", "англ. термин"),
        (r"₽/m³|(?<![а-яА-ЯёЁ])m³", "homoglyph m³"),
    ]
    for pat, label in patterns:
        if re.search(pat, text, re.I):
            issues.append(label)
    return issues


def main() -> None:
    parser = argparse.ArgumentParser(description="Normalize a DOCX file to the Russian business document standard")
    parser.add_argument("files", nargs="+", help="DOCX files to normalize")
    parser.add_argument("--size", type=int, default=12, help="Body font size pt")
    parser.add_argument("--flat", action="store_true", help="All text 12 pt including headings")
    parser.add_argument("--check-only", action="store_true", help="Only scan, do not modify")
    args = parser.parse_args()

    exit_code = 0
    for f in args.files:
        path = Path(f)
        if not path.exists():
            print(f"MISSING: {path}", file=sys.stderr)
            exit_code = 1
            continue
        issues = scan_bad_patterns(path)
        if args.check_only:
            if issues:
                print(f"FAIL {path.name}: {', '.join(issues)}")
                exit_code = 1
            else:
                print(f"OK {path.name}")
            continue
        process_docx(path, args.size, flat=args.flat)
        issues = scan_bad_patterns(path)
        if issues:
            print(f"WARN {path.name} after normalize: {', '.join(issues)}")
            exit_code = 1
        else:
            print(f"OK {path.name}")
    sys.exit(exit_code)


if __name__ == "__main__":
    main()
