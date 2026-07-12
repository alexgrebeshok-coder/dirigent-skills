#!/usr/bin/env python3
"""Format analytical DOCX: bordered tables, title block, no TOC."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HEADER_FILL = "1F3864"
LABEL_FILL = "F2F2F2"
TOTAL_FILL = "D6E4F0"
POS_FILL = "E2EFDA"
WARN_FILL = "FFF2CC"
NEG_FILL = "FCE4D6"
WHITE_FILL = "FFFFFF"

# These pattern lists are a starting point — tune the actual words/numbers to your own
# project's vocabulary before running this on real documents. Cell text is classified by
# simple substring/regex match, not real NLP, so false positives on generic words like
# "нет" are expected and worth reviewing by eye once per project.
NEG_PATTERNS = (
    r"не рекомен",
    r"дисквал",
    r"отриц",
    r"убыт",
    r"−",
    r"критич",
    r"не подтвержден",
    r"\bнет\b",
    r"^нет$",
    r"^нет,",
    r"^нет ",
    r"^нет\.",
)
WARN_PATTERNS = (
    r"н/д",
    r"нет данных",
    r"затруднена",
    r"средняя",
    r"⚠",
)
POS_PATTERNS = (
    r"\bда\b",
    r"высокая",
    r"tier",
    r"прибыл",
    r"рекоменд",
    r"член",
    r"вхождение",
)


def set_cell_borders(cell) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tc_pr.append(borders)


def set_cell_shading(cell, fill: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tbl_pr.append(borders)


def format_run(run, *, size: int = 12, bold: bool = False, white: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(255, 255, 255) if white else RGBColor(0, 0, 0)


def format_paragraph_runs(para, *, size: int = 12, bold: bool = False, white: bool = False) -> None:
    for run in para.runs:
        format_run(run, size=size, bold=bold, white=white)


def classify_fill(text: str, *, col_idx: int, is_high_energy_col: bool) -> str:
    t = (text or "").strip().lower()
    if not t:
        return WHITE_FILL
    if any(re.search(p, t) for p in WARN_PATTERNS):
        return WARN_FILL
    if any(re.search(p, t) for p in NEG_PATTERNS):
        return NEG_FILL
    if any(re.search(p, t) for p in POS_PATTERNS):
        return POS_FILL
    if re.fullmatch(r"[789]([.,]\d+)?", t.replace(" ", "")):
        return POS_FILL
    return WHITE_FILL


def is_total_row(cells_text: list[str]) -> bool:
    joined = " ".join(cells_text).lower()
    return "итого" in joined


def format_table(table) -> None:
    set_table_borders(table)
    rows = table.rows
    if not rows:
        return

    for ri, row in enumerate(rows):
        cells_text = [cell.text.strip() for cell in row.cells]
        total_row = is_total_row(cells_text)
        for ci, cell in enumerate(row.cells):
            set_cell_borders(cell)
            text = cells_text[ci]

            if ri == 0:
                fill = HEADER_FILL
                bold = True
                white = True
            elif total_row:
                fill = TOTAL_FILL if ci <= 1 else classify_fill(text, col_idx=ci, is_high_energy_col=ci == 1)
                bold = ci <= 1 or ("итого" in text.lower())
                white = False
            elif ci == 0:
                fill = LABEL_FILL
                bold = True
                white = False
            else:
                fill = classify_fill(text, col_idx=ci, is_high_energy_col=ci == 1)
                bold = False
                white = False

            set_cell_shading(cell, fill)
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                para.paragraph_format.line_spacing = 1.15
                if ri == 0:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                format_paragraph_runs(para, size=11, bold=bold, white=white)


def strip_unwanted_paragraphs(doc: Document) -> None:
    remove: list = []
    seen_main_title = False
    for para in doc.paragraphs:
        style = getattr(para.style, "name", "") or ""
        text = para.text.strip()

        if style in {"TOC Heading", "TOC"} or text == "Оглавление":
            remove.append(para)
            continue

        if style in {"Title", "Subtitle", "Date", "Author"}:
            remove.append(para)
            continue

        if style == "Heading 1" and not seen_main_title and text.lower().startswith("сравнительный"):
            remove.append(para)
            continue
        if style == "Heading 1" and not seen_main_title and text.lower().startswith("исследование"):
            remove.append(para)
            continue
        if style == "Heading 1" and not seen_main_title and text.lower().startswith("пояснительная"):
            remove.append(para)
            continue
        if style == "Heading 1" and not seen_main_title and text.lower().startswith("профиль"):
            remove.append(para)
            continue
        if style == "Heading 1" and not seen_main_title and text.lower().startswith("объединённый"):
            remove.append(para)
            continue

        if style == "Heading 1" and re.match(r"^\d+\.", text):
            seen_main_title = True

    for para in remove:
        p = para._element
        p.getparent().remove(p)


def format_title_block(doc: Document, *, doc_kind: str) -> None:
    body = doc.element.body
    paras = doc.paragraphs
    if not paras:
        return

    # Placeholder title/subtitle strings — replace "Example Partner LLC" and the subject
    # line with your own project's wording before running this on a real document.
    if doc_kind == "compare":
        title = "СРАВНИТЕЛЬНЫЙ АНАЛИЗ ПАРТНЁРОВ"
        subtitle = "Пример: сравнение нескольких компаний по заданным критериям"
    elif doc_kind == "partner_internal":
        title = "ПОЯСНИТЕЛЬНАЯ ЗАПИСКА (ВНУТРЕННЯЯ)"
        subtitle = "Партнёрские возможности ООО «Пример-Партнёр» — внутренняя оценка"
    elif doc_kind == "partner_external":
        title = "ПРОФИЛЬ ПАРТНЁРА"
        subtitle = "ООО «Пример-Партнёр» — возможности сотрудничества"
    elif doc_kind == "merged_4":
        title = "ОБЪЕДИНЁННЫЙ АНАЛИТИЧЕСКИЙ ОТЧЁТ"
        subtitle = "Несколько компаний | Сравнение и углублённая проверка контрагентов"
    else:
        title = "ИССЛЕДОВАНИЕ ПАРТНЁРОВ И РЫНКА"
        subtitle = "Проверка контрагентов, карта рынка и краткий список партнёров"

    first = paras[0]
    first.text = title
    first.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first.paragraph_format.space_after = Pt(6)
    format_paragraph_runs(first, size=16, bold=True)

    if len(paras) > 1 and paras[1].text.strip():
        sub = paras[1]
    else:
        sub = doc.add_paragraph()
        body.insert(body.index(first._element) + 1, sub._element)
    sub.text = subtitle
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(12)
    format_paragraph_runs(sub, size=12, bold=False)


def format_document_paragraphs(doc: Document) -> None:
    for para in doc.paragraphs:
        style = getattr(para.style, "name", "") or ""
        text = para.text.strip()
        if not text:
            continue

        if style.startswith("Heading 1") or (style == "Heading 2" and re.match(r"^\d+\.", text)):
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            format_paragraph_runs(para, size=14, bold=True)
            continue

        if style.startswith("Heading"):
            format_paragraph_runs(para, size=12, bold=True)
            continue

        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para.paragraph_format.line_spacing = 1.15
        format_paragraph_runs(para, size=12, bold=False)


def set_page_layout(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def process_docx(path: Path, doc_kind: str) -> None:
    doc = Document(str(path))
    strip_unwanted_paragraphs(doc)
    format_title_block(doc, doc_kind=doc_kind)
    format_document_paragraphs(doc)
    for table in doc.tables:
        format_table(table)
    set_page_layout(doc)
    doc.save(str(path))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("files", nargs="+")
    parser.add_argument(
        "--kind",
        choices=("compare", "research", "partner_internal", "partner_external", "merged_4"),
        default="compare",
    )
    args = parser.parse_args()
    for f in args.files:
        process_docx(Path(f), doc_kind=args.kind)
        print(f"OK {Path(f).name}")


if __name__ == "__main__":
    main()
